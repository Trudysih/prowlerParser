import pandas as pd
import re
import time
import os
import pprint
import docx
from docx.shared import Cm, Inches

# Helper function to extract check number from 'TITLE_TEXT' column
# Input: 'TITLE_TEXT' column string
# Output: String within square parentheses'[]', assuming the check number is always the first occurence


def getCheckNumberWithText(string):
    originalString = string
    newString = re.search(r"\[(\w+)\]", originalString)
    return newString.group(1)


def getCheckNumber(string):
    s = getCheckNumberWithText(string)
    return int(re.search(r'\d+', s).group())


def getCheckDescription(string):
    o_string = string
    pattern = r'\[[^\]]*\] '
    pattern2 = r' \(.*\)'
    t = re.sub(pattern, '', o_string)
    t2 = re.sub(pattern2, '', t)
    return t2


def set_col_widths(table):
    widths = (Inches(0.5), Inches(6.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def write_to_word_doc(dictChecks):
    doc = docx.Document()
    doc.add_heading('Automation report', 0)

    # loop thru the dictionary, create 1 table for each object in the dictionary
    for k, v in dictChecks.items():
        sample_df = v
        tableDescription = str(k) + ": " + str(sample_df.values[0][3])
        # create table
        t = doc.add_table(sample_df.shape[0]+1, cols=1)
        t.style = 'Table Grid'

        # add the header row + make it bold
        t.cell(0, 0).paragraphs[0].add_run(tableDescription + '\n').bold = True

        # add the rest of the data frame
        for i in range(sample_df.shape[0]):  # number of rows
            t.cell(i+1, 0).text = str(sample_df.values[i, 2])

        # add spacing between each table
        doc.add_paragraph('\n')
        doc.save(output_file_word)


# Read in CSV file from prowler
data = pd.read_csv('output.csv')

# Filter results by columns I want
data_filtered = data[['REGION', 'RESULT', 'TITLE_TEXT', 'NOTES']]

# Select all where RESULT == FAIL
df = (data_filtered[data_filtered.RESULT == 'FAIL'])
df = df[['REGION', 'TITLE_TEXT', 'NOTES']]

# Extract checkNumber from 'TITLE_TEXT' field
df['CHECK_DESCRIPTION'] = df['TITLE_TEXT'].apply(getCheckDescription)
df['TITLE_TEXT'] = df['TITLE_TEXT'].apply(getCheckNumber)

# rename 'TITLE_TEXT' to 'CHECK_NUMBER'
df.rename(columns={'TITLE_TEXT': 'CHECK_NUMBER'}, inplace=True)

# Write to output CSV file
t = time.localtime()
timestamp = time.strftime('%H%M%S_%d%m%Y', t)
path = 'output'
output_file_csv = os.path.join(path, 'output_'+timestamp+'.csv')
output_file_word = os.path.join(path, 'output_'+timestamp+'.docx')

# create dictionary of the dataframes that u want
dictionary_of_checks = {k: v for k, v in df.groupby('CHECK_NUMBER')}
# Use the dictionary to write to docx document
write_to_word_doc(dictionary_of_checks)
